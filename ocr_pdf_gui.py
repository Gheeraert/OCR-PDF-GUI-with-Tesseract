#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
OCR PDF → TXT / DOCX (+ PDF interrogeable), gestion des doubles pages, détection gouttière améliorée
- Gouttière : seuillage Otsu + projection verticale lissée
- Fenêtre centrale ajustable + lissage ajustable
- Bouton d'aperçu de la coupe sur la première page
"""

import io
import os
import sys
import threading
import queue
import shutil
import re
from datetime import datetime

import tkinter as tk
from tkinter import filedialog, messagebox, ttk

try:
    import fitz  # PyMuPDF
except Exception as e:
    raise SystemExit("PyMuPDF (fitz) requis. Installez-le via 'pip install pymupdf'.") from e

try:
    from PIL import Image, ImageOps, ImageFilter, ImageTk, ImageDraw
except Exception as e:
    raise SystemExit("Pillow requis. Installez-le via 'pip install pillow'.") from e

try:
    import pytesseract
except Exception as e:
    raise SystemExit("pytesseract requis. Installez-le via 'pip install pytesseract'.") from e

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None


# ====== Utilitaires ======

def configure_tesseract(explicit_path=None):
    candidates = []
    if explicit_path:
        candidates.append(explicit_path)
    which = shutil.which("tesseract")
    if which:
        candidates.append(which)
    candidates += [
        r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
        r"C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe",
    ]
    seen = set(); ordered = []
    for c in candidates:
        if c and c not in seen:
            seen.add(c); ordered.append(c)
    for c in ordered:
        if os.path.isfile(c):
            pytesseract.pytesseract.tesseract_cmd = c
            return c
    return None


def pil_from_fitz_pixmap(pix):
    data = pix.tobytes("png")
    return Image.open(io.BytesIO(data)).convert("RGB")


def otsu_thresh(arr):
    import numpy as np
    hist, _ = np.histogram(arr.flatten(), bins=256, range=(0,255))
    total = arr.size
    sum_total = (hist * np.arange(256)).sum()
    sumB = 0.0; wB = 0.0; varMax = 0.0; thresh = 0
    for t in range(256):
        wB += hist[t]
        if wB == 0:
            continue
        wF = total - wB
        if wF == 0:
            break
        sumB += t * hist[t]
        mB = sumB / wB
        mF = (sum_total - sumB) / wF
        var_between = wB * wF * (mB - mF) ** 2
        if var_between > varMax:
            varMax = var_between
            thresh = t
    return thresh


def find_gutter_x(img, central_frac=0.4, smooth_px=25):
    """Détecte la gouttière près du centre par projection verticale lissée.
    central_frac : proportion du centre utilisée (0.2-0.8)
    smooth_px    : largeur de lissage (en pixels) sur l'image de travail
    """
    import numpy as np
    g = img.convert("L")
    g = ImageOps.autocontrast(g)
    g = g.filter(ImageFilter.MedianFilter(3))
    a = np.array(g, dtype=np.uint8)
    t = otsu_thresh(a)
    bw = (a < t).astype(np.uint8)  # 1 = encre, 0 = fond
    h, w = bw.shape
    top = int(h * 0.05); bottom = int(h * 0.95)
    work = bw[top:bottom, :]
    col = work.sum(axis=0).astype(float)
    # lissage par convolution
    k = max(3, int(smooth_px) | 1)  # impair
    kern = np.ones(k, dtype=float) / k
    sm = np.convolve(col, kern, mode="same")
    # recherche dans une bande centrale
    cf = float(central_frac)
    cf = min(0.9, max(0.1, cf))
    left = int(w * (0.5 - cf/2))
    right = int(w * (0.5 + cf/2))
    idx = sm[left:right].argmin()
    cut_x = left + int(idx)
    margin = max(10, int(w * 0.01))
    cut_x = max(margin, min(w - margin, cut_x))
    return cut_x


def split_double_page(img, mode="auto", central_frac=0.4, smooth_px=25):
    w, h = img.size
    if mode == "half":
        gx = w // 2
    else:
        # Heuristique : seulement si largeur/hauteur >= 1.28
        if (w / h) < 1.28:
            return [img]
        gx = find_gutter_x(img, central_frac=central_frac, smooth_px=smooth_px)
    left = img.crop((0, 0, gx, h))
    right = img.crop((gx, 0, w, h))
    return [left, right]


IMAGE_EXTENSIONS = (".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp")


def natural_sort_key(s):
    return [int(t) if t.isdigit() else t.lower() for t in re.split(r"(\d+)", s)]


def list_images_in_folder(folder):
    names = [n for n in os.listdir(folder) if n.lower().endswith(IMAGE_EXTENSIONS)]
    names.sort(key=natural_sort_key)
    return [os.path.join(folder, n) for n in names]


def normalize_illumination(gray_img, blur_radius=31):
    """Corrige un éclairage non uniforme (ombre portée, gradient de fenêtre)
    en divisant l'image par une version très floutée d'elle-même (flat-fielding)."""
    import numpy as np
    arr = np.asarray(gray_img, dtype=np.float32)
    bg = np.asarray(gray_img.filter(ImageFilter.GaussianBlur(radius=blur_radius)), dtype=np.float32)
    bg[bg < 1.0] = 1.0
    norm = np.clip(arr / bg * 255.0, 0, 255).astype(np.uint8)
    return Image.fromarray(norm, mode="L")


def adaptive_threshold(gray_img, blur_radius=15, offset=10):
    """Seuillage adaptatif (comparaison à une moyenne locale) : plus robuste qu'un
    seuil Otsu global face à un éclairage encore légèrement inégal après normalisation."""
    import numpy as np
    arr = np.asarray(gray_img, dtype=np.float32)
    local_mean = np.asarray(gray_img.filter(ImageFilter.BoxBlur(blur_radius)), dtype=np.float32)
    binary = np.where(arr > (local_mean - offset), 255, 0).astype(np.uint8)
    return Image.fromarray(binary, mode="L")


def estimate_skew_angle(gray_img, max_angle=5.0, step=0.5):
    """Cherche l'angle (dans [-max_angle, +max_angle]) qui maximise la variance des
    projections horizontales — plus l'image est droite, plus les lignes de texte
    forment des bandes nettes (forte variance de la somme d'encre par ligne)."""
    import numpy as np
    small = gray_img
    max_dim = 1000
    if max(small.size) > max_dim:
        scale = max_dim / max(small.size)
        small = small.resize((max(1, int(small.width * scale)), max(1, int(small.height * scale))))
    arr = np.asarray(small, dtype=np.float32)
    bw = (arr < arr.mean()).astype(np.uint8) * 255
    bw_img = Image.fromarray(bw, mode="L")

    best_angle, best_score = 0.0, -1.0
    angle = -abs(max_angle)
    while angle <= abs(max_angle) + 1e-9:
        rotated = bw_img.rotate(angle, resample=Image.BICUBIC, expand=False, fillcolor=0)
        row_sums = np.asarray(rotated, dtype=np.float32).sum(axis=1)
        score = float(row_sums.var())
        if score > best_score:
            best_score, best_angle = score, angle
        angle += step
    return best_angle


def deskew_image(img, max_angle=5.0, step=0.5):
    """Redresse une image inclinée de quelques degrés (photo prise à main levée)."""
    gray = img.convert("L") if img.mode != "L" else img
    angle = estimate_skew_angle(gray, max_angle=max_angle, step=step)
    if abs(angle) < 0.05:
        return img, 0.0
    fill = 255 if img.mode == "L" else (255, 255, 255)
    rotated = img.rotate(angle, resample=Image.BICUBIC, expand=True, fillcolor=fill)
    return rotated, angle


def postprocess_text(txt: str) -> str:
    txt = txt.replace("\r\n", "\n").replace("\r", "\n")
    txt = re.sub(r"(\w)[\-­]\n(\w)", r"\1\2", txt)
    txt = re.sub(r"(?<!\n)\n(?!\n)", " ", txt)
    txt = re.sub(r"[ \t]{2,}", " ", txt)
    return txt.strip()


# ====== Application ======

class OCRApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("OCR PDF → TXT / DOCX (+ PDF interrogeable)")
        self.geometry("880x700")
        self.minsize(860, 680)

        self.input_mode_var = tk.StringVar(value="pdf")  # "pdf" ou "images"
        self.pdf_path_var = tk.StringVar()
        self.image_dir_var = tk.StringVar()
        self.out_dir_var = tk.StringVar()
        self.lang_var = tk.StringVar(value="fra+eng")
        self.res_dpi_var = tk.IntVar(value=300)
        self.format_txt_var = tk.BooleanVar(value=True)
        self.format_docx_var = tk.BooleanVar(value=True if Document is not None else False)
        self.format_pdf_var = tk.BooleanVar(value=False)
        self.keep_pagebreaks_var = tk.BooleanVar(value=True)
        self.page_range_var = tk.StringVar(value="all")
        self.tess_config_var = tk.StringVar(value="--oem 1 --psm 6")

        # Doublage
        self.split_doubles_var = tk.BooleanVar(value=True)
        self.split_mode_var = tk.StringVar(value="auto")
        self.central_frac_var = tk.DoubleVar(value=0.4)   # 40% du centre
        self.smooth_px_var = tk.IntVar(value=25)

        # Prétraitement "photo smartphone"
        self.photo_mode_var = tk.BooleanVar(value=False)
        self.max_skew_var = tk.DoubleVar(value=5.0)

        self._worker = None
        self._cancel_event = threading.Event()
        self._ui_queue = queue.Queue()
        self._build_ui()
        self._on_input_mode_change()
        self._poll_queue()

    def _build_ui(self):
        pad = {'padx': 10, 'pady': 6}
        frm = ttk.Frame(self); frm.pack(fill="both", expand=True)

        # Source d'entrée : PDF ou dossier de photos
        row0 = ttk.Frame(frm); row0.pack(fill="x", **pad)
        ttk.Label(row0, text="Source :").pack(side="left")
        ttk.Radiobutton(row0, text="Fichier PDF", variable=self.input_mode_var, value="pdf",
                         command=self._on_input_mode_change).pack(side="left", padx=(8, 4))
        ttk.Radiobutton(row0, text="Dossier de photos (JPG/PNG…)", variable=self.input_mode_var, value="images",
                         command=self._on_input_mode_change).pack(side="left", padx=4)

        row1 = ttk.Frame(frm); row1.pack(fill="x", **pad)
        self.pdf_label = ttk.Label(row1, text="Fichier PDF :"); self.pdf_label.pack(side="left")
        self.pdf_entry = ttk.Entry(row1, textvariable=self.pdf_path_var)
        self.pdf_entry.pack(side="left", fill="x", expand=True, padx=8)
        self.pdf_browse_btn = ttk.Button(row1, text="Parcourir…", command=self._browse_pdf)
        self.pdf_browse_btn.pack(side="left")

        row1b = ttk.Frame(frm); row1b.pack(fill="x", **pad)
        self.imgdir_label = ttk.Label(row1b, text="Dossier de photos :"); self.imgdir_label.pack(side="left")
        self.imgdir_entry = ttk.Entry(row1b, textvariable=self.image_dir_var)
        self.imgdir_entry.pack(side="left", fill="x", expand=True, padx=8)
        self.imgdir_browse_btn = ttk.Button(row1b, text="Parcourir…", command=self._browse_image_dir)
        self.imgdir_browse_btn.pack(side="left")

        # Dossier sortie
        row2 = ttk.Frame(frm); row2.pack(fill="x", **pad)
        ttk.Label(row2, text="Dossier de sortie :").pack(side="left")
        ttk.Entry(row2, textvariable=self.out_dir_var).pack(side="left", fill="x", expand=True, padx=8)
        ttk.Button(row2, text="Choisir…", command=self._browse_outdir).pack(side="left")

        # Paramètres OCR
        row3 = ttk.LabelFrame(frm, text="Paramètres OCR"); row3.pack(fill="x", **pad)
        ttk.Label(row3, text="Langues (Tesseract):").grid(row=0, column=0, sticky="w", padx=8, pady=4)
        ttk.Entry(row3, textvariable=self.lang_var, width=20).grid(row=0, column=1, sticky="w")
        ttk.Label(row3, text="Config Tesseract :").grid(row=0, column=2, sticky="e", padx=8)
        ttk.Entry(row3, textvariable=self.tess_config_var, width=28).grid(row=0, column=3, sticky="w")
        ttk.Label(row3, text="Résolution (DPI) :").grid(row=1, column=0, sticky="w", padx=8, pady=4)
        ttk.Spinbox(row3, from_=150, to=600, increment=25, textvariable=self.res_dpi_var, width=7).grid(row=1, column=1, sticky="w")
        ttk.Label(row3, text="Pages (ex. 'all', '1-5', '1-3,6,8-') :").grid(row=1, column=2, sticky="e", padx=8)
        ttk.Entry(row3, textvariable=self.page_range_var, width=28).grid(row=1, column=3, sticky="w")
        ttk.Checkbutton(row3, text="Insérer des sauts de page", variable=self.keep_pagebreaks_var).grid(row=2, column=0, columnspan=2, sticky="w", padx=8, pady=4)

        # Formats
        row4 = ttk.LabelFrame(frm, text="Format de sortie"); row4.pack(fill="x", **pad)
        ttk.Checkbutton(row4, text="Texte (.txt)", variable=self.format_txt_var).pack(side="left", padx=8, pady=4)
        ttk.Checkbutton(row4, text="Word (.docx)", variable=self.format_docx_var, state=("normal" if Document is not None else "disabled")).pack(side="left", padx=8, pady=4)
        ttk.Checkbutton(row4, text="PDF interrogeable (.pdf)", variable=self.format_pdf_var).pack(side="left", padx=8, pady=4)

        # Doubles pages
        row4b = ttk.LabelFrame(frm, text="Doubles pages"); row4b.pack(fill="x", **pad)
        ttk.Checkbutton(row4b, text="Couper les doubles pages", variable=self.split_doubles_var).grid(row=0, column=0, sticky="w", padx=8, pady=4)
        ttk.Label(row4b, text="Méthode :").grid(row=0, column=1, sticky="e")
        ttk.Combobox(row4b, textvariable=self.split_mode_var, width=10, state="readonly", values=["auto", "half"]).grid(row=0, column=2, sticky="w")
        ttk.Label(row4b, text="Fenêtre centrale (0.2–0.8) :").grid(row=1, column=1, sticky="e")
        ttk.Spinbox(row4b, from_=0.2, to=0.8, increment=0.05, textvariable=self.central_frac_var, width=6).grid(row=1, column=2, sticky="w")
        ttk.Label(row4b, text="Lissage (px) :").grid(row=1, column=3, sticky="e")
        ttk.Spinbox(row4b, from_=5, to=101, increment=2, textvariable=self.smooth_px_var, width=6).grid(row=1, column=4, sticky="w")
        ttk.Button(row4b, text="Aperçu découpe (p.1)", command=self.preview_cut).grid(row=0, column=3, columnspan=2, padx=12)

        # Prétraitement photo smartphone
        row4c = ttk.LabelFrame(frm, text="Prétraitement photo (smartphone)"); row4c.pack(fill="x", **pad)
        ttk.Checkbutton(
            row4c, text="Mode photo (correction d’éclairage + redressement + seuillage adaptatif)",
            variable=self.photo_mode_var,
        ).grid(row=0, column=0, columnspan=3, sticky="w", padx=8, pady=4)
        ttk.Label(row4c, text="Inclinaison max. corrigée (°) :").grid(row=1, column=0, sticky="w", padx=8)
        ttk.Spinbox(row4c, from_=1, to=15, increment=1, textvariable=self.max_skew_var, width=6).grid(row=1, column=1, sticky="w")
        ttk.Label(
            row4c,
            text="Recommandé pour des photos prises à main levée (à la place du simple contraste/médiane).",
        ).grid(row=2, column=0, columnspan=3, sticky="w", padx=8, pady=(0, 4))

        # Progression
        row5 = ttk.Frame(frm); row5.pack(fill="x", **pad)
        self.progress = ttk.Progressbar(row5, mode="determinate"); self.progress.pack(fill="x", padx=4)
        self.status_var = tk.StringVar(value="Prêt.")
        ttk.Label(row5, textvariable=self.status_var).pack(anchor="w", padx=4, pady=2)

        # Boutons
        row6 = ttk.Frame(frm); row6.pack(fill="x", **pad)
        ttk.Button(row6, text="Démarrer l’OCR", command=self.start_ocr).pack(side="left", padx=4)
        ttk.Button(row6, text="Annuler", command=self.cancel_ocr).pack(side="left", padx=4)
        ttk.Button(row6, text="Ouvrir le dossier", command=self._open_outdir).pack(side="left", padx=4)
        ttk.Button(row6, text="Quitter", command=self.destroy).pack(side="right", padx=4)

        # Journal
        row7 = ttk.LabelFrame(frm, text="Journal"); row7.pack(fill="both", expand=True, **pad)
        self.log = tk.Text(row7, height=12, wrap="word"); self.log.pack(fill="both", expand=True, padx=6, pady=6)

    def logln(self, text):
        self.log.insert("end", text + "\n"); self.log.see("end"); self.update_idletasks()

    def set_status(self, text):
        self.status_var.set(text); self.update_idletasks()

    # ---- Communication thread-safe avec le worker OCR ----
    def _q_put(self, kind, payload=None):
        """Appelé depuis le thread worker : dépose un message pour le thread UI."""
        self._ui_queue.put((kind, payload))

    def _poll_queue(self):
        """Tourne sur le thread principal (via .after) et applique les messages du worker."""
        try:
            while True:
                kind, payload = self._ui_queue.get_nowait()
                if kind == "log":
                    self.logln(payload)
                elif kind == "status":
                    self.set_status(payload)
                elif kind == "progress_max":
                    self.progress["maximum"] = payload
                elif kind == "progress_value":
                    self.progress["value"] = payload
                elif kind == "info":
                    messagebox.showinfo(*payload)
                elif kind == "error":
                    messagebox.showerror(*payload)
        except queue.Empty:
            pass
        self.after(100, self._poll_queue)

    def cancel_ocr(self):
        if self._worker and self._worker.is_alive():
            self._cancel_event.set()
            self.set_status("Annulation demandée…")
        else:
            self.set_status("Aucun traitement en cours.")

    def _on_input_mode_change(self):
        images_mode = self.input_mode_var.get() == "images"
        state_pdf = "disabled" if images_mode else "normal"
        state_img = "normal" if images_mode else "disabled"
        self.pdf_entry.configure(state=state_pdf)
        self.pdf_browse_btn.configure(state=state_pdf)
        self.imgdir_entry.configure(state=state_img)
        self.imgdir_browse_btn.configure(state=state_img)
        if images_mode:
            self.photo_mode_var.set(True)

    def _browse_pdf(self):
        path = filedialog.askopenfilename(title="Choisir un PDF", filetypes=[("PDF", "*.pdf"), ("Tous les fichiers", "*.*")])
        if path:
            self.pdf_path_var.set(path); self.out_dir_var.set(os.path.dirname(path))

    def _browse_image_dir(self):
        d = filedialog.askdirectory(title="Choisir un dossier de photos")
        if d:
            self.image_dir_var.set(d); self.out_dir_var.set(d)

    def _browse_outdir(self):
        d = filedialog.askdirectory(title="Choisir un dossier de sortie")
        if d: self.out_dir_var.set(d)

    def _open_outdir(self):
        d = self.out_dir_var.get().strip()
        if not d:
            messagebox.showinfo("Dossier", "Aucun dossier de sortie défini."); return
        try:
            if sys.platform.startswith('win'): os.startfile(d)
            elif sys.platform == 'darwin': os.system(f'open "{d}"')
            else: os.system(f'xdg-open "{d}"')
        except Exception as e:
            messagebox.showerror("Erreur", f"Impossible d’ouvrir le dossier:\n{e}")

    def _parse_page_ranges(self, pr_str, nb_pages):
        pr_str = pr_str.strip().lower()
        if pr_str in ("all", "", "*"): return list(range(nb_pages))
        indices = set()
        for p in [x.strip() for x in pr_str.split(",") if x.strip()]:
            if "-" in p:
                a, b = p.split("-", 1); a=a.strip(); b=b.strip()
                start = int(a) if a else 1; end = int(b) if b else nb_pages
                for i in range(max(1,start)-1, min(nb_pages,end)): indices.add(i)
            else:
                i = int(p)-1
                if 0 <= i < nb_pages: indices.add(i)
        return sorted(indices)

    def _load_first_page_image(self):
        """Retourne la première page/photo sous forme d'image PIL, selon la source active."""
        if self.input_mode_var.get() == "images":
            image_dir = self.image_dir_var.get().strip()
            if not image_dir or not os.path.isdir(image_dir):
                raise ValueError("Choisis d'abord un dossier de photos.")
            paths = list_images_in_folder(image_dir)
            if not paths:
                raise ValueError("Aucune image (JPG/PNG/…) trouvée dans ce dossier.")
            return Image.open(paths[0]).convert("RGB")
        else:
            pdf_path = self.pdf_path_var.get().strip()
            if not pdf_path or not os.path.isfile(pdf_path):
                raise ValueError("Choisis d'abord un PDF.")
            doc = fitz.open(pdf_path)
            page = doc.load_page(0)
            dpi = int(self.res_dpi_var.get())
            zoom = dpi / 72.0; mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            return pil_from_fitz_pixmap(pix)

    def preview_cut(self):
        try:
            img = self._load_first_page_image()
            if self.split_mode_var.get() == "half":
                gx = img.width // 2
            else:
                gx = find_gutter_x(img, self.central_frac_var.get(), self.smooth_px_var.get())
            prev = img.copy()
            draw = ImageDraw.Draw(prev)
            draw.line([(gx,0),(gx,prev.height)], fill=(255,0,0), width=6)
            # scale preview
            max_w = 1200
            scale = min(1.0, max_w / prev.width)
            show = prev.resize((int(prev.width*scale), int(prev.height*scale)))
            top = tk.Toplevel(self); top.title("Aperçu découpe (p.1)")
            tkimg = ImageTk.PhotoImage(show)
            lbl = ttk.Label(top, image=tkimg); lbl.image = tkimg
            lbl.pack()
        except Exception as e:
            messagebox.showerror("Erreur", str(e))

    def start_ocr(self):
        if self._worker and self._worker.is_alive():
            messagebox.showinfo("En cours", "Un traitement est déjà en cours."); return
        images_mode = self.input_mode_var.get() == "images"
        pdf_path = self.pdf_path_var.get().strip()
        image_dir = self.image_dir_var.get().strip()
        out_dir = self.out_dir_var.get().strip()
        if images_mode:
            if not image_dir or not os.path.isdir(image_dir):
                messagebox.showerror("Erreur", "Veuillez choisir un dossier de photos valide."); return
            if not list_images_in_folder(image_dir):
                messagebox.showerror("Erreur", "Aucune image (JPG/PNG/…) trouvée dans ce dossier."); return
        elif not pdf_path or not os.path.isfile(pdf_path):
            messagebox.showerror("Erreur", "Veuillez choisir un PDF valide."); return
        if not out_dir or not os.path.isdir(out_dir):
            messagebox.showerror("Erreur", "Veuillez choisir un dossier de sortie valide."); return
        if not (self.format_txt_var.get() or self.format_docx_var.get() or self.format_pdf_var.get()):
            messagebox.showerror("Erreur", "Sélectionnez au moins un format (TXT, DOCX ou PDF)."); return
        if self.format_docx_var.get() and Document is None:
            messagebox.showerror("Erreur", "python-docx n'est pas installé. 'pip install python-docx' ou décochez DOCX."); return
        found = configure_tesseract()
        if not found:
            messagebox.showerror("Tesseract introuvable", "Je ne trouve pas tesseract.exe. Vérifie l’installation / PATH."); return

        try:
            langs_avail = pytesseract.get_languages(config="")
            self.logln(f"Langues Tesseract disponibles: {', '.join(sorted(langs_avail))}")
        except Exception:
            self.logln("Impossible d’énumérer les langues Tesseract." )

        self.log.delete("1.0", "end"); self.progress["value"] = 0
        self.set_status("Préparation…")
        if images_mode:
            self.logln(f"Dossier de photos : {image_dir}")
        else:
            self.logln(f"PDF : {pdf_path}")
        self.logln(f"Dossier de sortie : {out_dir}"); self.logln(f"Tesseract : {found}")
        self._cancel_event.clear()
        self._worker = threading.Thread(target=self._do_ocr, daemon=True); self._worker.start()

    def _do_ocr(self):
        # Tourne dans un thread worker : toute communication avec l'UI passe
        # par self._q_put (jamais d'accès direct aux widgets depuis ce thread).
        txt_file = None
        docx = None
        final_pdf = None
        cancelled = False
        SAVE_EVERY = 10  # sauvegarde périodique DOCX/PDF pour ne rien perdre en cas de crash/annulation

        try:
            images_mode = self.input_mode_var.get() == "images"
            pdf_path = self.pdf_path_var.get().strip()
            image_dir = self.image_dir_var.get().strip()
            out_dir = self.out_dir_var.get().strip()
            dpi = int(self.res_dpi_var.get())
            langs = self.lang_var.get().strip() or "fra"
            pr_str = self.page_range_var.get().strip()
            tess_cfg = self.tess_config_var.get().strip()
            photo_mode = bool(self.photo_mode_var.get())
            max_skew = float(self.max_skew_var.get())

            if images_mode:
                image_paths = list_images_in_folder(image_dir)
                nb_pages = len(image_paths)
                source_label = os.path.basename(image_dir.rstrip("\\/")) or image_dir
                doc = None
            else:
                doc = fitz.open(pdf_path)
                nb_pages = doc.page_count
                source_label = os.path.basename(pdf_path)

            def load_page_image(idx):
                if images_mode:
                    return Image.open(image_paths[idx]).convert("RGB")
                page = doc.load_page(idx)
                zoom = dpi / 72.0; mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                return pil_from_fitz_pixmap(pix)

            base = os.path.splitext(source_label)[0] if not images_mode else source_label
            txt_out = os.path.join(out_dir, f"{base}_OCR.txt")
            docx_out = os.path.join(out_dir, f"{base}_OCR.docx")
            pdf_out = os.path.join(out_dir, f"{base}_OCR.pdf")
            self._q_put("log", f"Langues Tesseract : {langs}")
            if images_mode:
                self._q_put("log", f"{nb_pages} photo(s) trouvée(s)")
            else:
                self._q_put("log", f"Résolution rendu : {dpi} DPI")
            self._q_put("log", f"Pages : {pr_str}")
            self._q_put("log", f"Config Tesseract : {tess_cfg or '(par défaut)'}")
            self._q_put("log", f"Mode photo (éclairage + redressement) : {'oui' if photo_mode else 'non'}")

            pages_idx = self._parse_page_ranges(pr_str, nb_pages)
            if not pages_idx: raise ValueError("Aucune page à traiter.")
            self._q_put("progress_max", len(pages_idx))
            self._q_put("status", "OCR en cours…")

            if self.format_txt_var.get():
                txt_file = open(txt_out, "w", encoding="utf-8")
            txt_first_entry = True

            if self.format_docx_var.get() and Document is not None:
                docx = Document()
                try:
                    style = docx.styles["Normal"]; font = style.font; font.name = "Calibri"; font.size = Pt(11)
                except Exception: pass
                docx.add_heading(f"OCR de {source_label}", level=1)
                docx.add_paragraph(f"Généré le {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            if self.format_pdf_var.get():
                final_pdf = fitz.open()

            suffixes = "abcdefghijklmnopqrstuvwxyz"
            processed_sub_pages = 0

            for k, i in enumerate(pages_idx, start=1):
                if self._cancel_event.is_set():
                    cancelled = True
                    self._q_put("log", "Traitement annulé par l'utilisateur.")
                    break

                img = load_page_image(i)

                images_to_ocr = [img]
                if self.split_doubles_var.get():
                    images_to_ocr = split_double_page(img, mode=self.split_mode_var.get(),
                                                      central_frac=self.central_frac_var.get(),
                                                      smooth_px=int(self.smooth_px_var.get()))

                for sub_idx, sub_img in enumerate(images_to_ocr):
                    if self._cancel_event.is_set():
                        cancelled = True
                        self._q_put("log", "Traitement annulé par l'utilisateur.")
                        break

                    proc = sub_img.convert("L")
                    if photo_mode:
                        deskewed, angle = deskew_image(proc, max_angle=max_skew)
                        if angle:
                            self._q_put("log", f"[Page {i+1}] Redressement : {angle:+.1f}°")
                        proc = normalize_illumination(deskewed)
                        proc = adaptive_threshold(proc)
                    else:
                        proc = ImageOps.autocontrast(proc); proc = proc.filter(ImageFilter.MedianFilter(3))
                    try:
                        txt = pytesseract.image_to_string(proc, lang=langs, config=tess_cfg or None)
                    except Exception as e:
                        self._q_put("log", f"[Page {i+1}] Erreur OCR ({langs}) : {e}. Tentative de repli en 'eng'.")
                        try:
                            txt = pytesseract.image_to_string(proc, lang="eng", config=tess_cfg or None)
                        except Exception as e2:
                            self._q_put("log", f"[Page {i+1}] Échec du repli 'eng' : {e2}. Page laissée vide, poursuite du traitement.")
                            txt = ""

                    txt = postprocess_text(txt)
                    suffix = suffixes[sub_idx] if len(images_to_ocr) > 1 else ""
                    page_header = f"\n===== Page {i+1}{suffix}/{nb_pages} =====\n"
                    entry = page_header + txt + "\n"
                    self._q_put("log", f"[Page {i+1}{suffix}] {len(txt)} caractères extraits")

                    if txt_file is not None:
                        sep = "\n\f\n" if self.keep_pagebreaks_var.get() else "\n\n"
                        if not txt_first_entry:
                            txt_file.write(sep)
                        txt_file.write(entry.strip())
                        txt_file.flush()
                        txt_first_entry = False

                    if docx is not None:
                        docx.add_heading(f"Page {i+1}{suffix}/{nb_pages}", level=2)
                        docx.add_paragraph(txt)
                        if self.keep_pagebreaks_var.get():
                            docx.add_page_break()

                    if final_pdf is not None:
                        try:
                            pdf_bytes = pytesseract.image_to_pdf_or_hocr(proc, lang=langs, config=tess_cfg or None, extension='pdf')
                            page_pdf = fitz.open(stream=pdf_bytes, filetype='pdf')
                            final_pdf.insert_pdf(page_pdf)
                        except Exception as e:
                            self._q_put("log", f"[Page {i+1}{suffix}] PDF interrogeable impossible : {e}")

                    processed_sub_pages += 1
                    if processed_sub_pages % SAVE_EVERY == 0:
                        if docx is not None:
                            try: docx.save(docx_out)
                            except Exception as e: self._q_put("log", f"Sauvegarde intermédiaire DOCX impossible : {e}")
                        if final_pdf is not None:
                            try: final_pdf.save(pdf_out)
                            except Exception as e: self._q_put("log", f"Sauvegarde intermédiaire PDF impossible : {e}")

                if cancelled:
                    break

                self._q_put("progress_value", k)
                self._q_put("status", f"Page {k} / {len(pages_idx)} traitée")

            if docx is not None:
                docx.save(docx_out)
                self._q_put("log", f"DOCX enregistré : {docx_out}")

            if final_pdf is not None:
                final_pdf.save(pdf_out)
                self._q_put("log", f"PDF interrogeable enregistré : {pdf_out}")

            if txt_file is not None:
                self._q_put("log", f"TXT enregistré : {txt_out}")

            if cancelled:
                self._q_put("status", "Annulé (résultats partiels enregistrés)")
                self._q_put("log", "Traitement interrompu : les résultats partiels ont été enregistrés.")
                self._q_put("info", ("OCR annulé", "Traitement annulé. Les pages déjà traitées ont été enregistrées."))
            else:
                self._q_put("status", "Terminé ✅")
                self._q_put("log", "Terminé.")
                self._q_put("info", ("OCR", "Traitement terminé."))

        except Exception as e:
            self._q_put("status", "Erreur ❌")
            self._q_put("log", f"ERREUR : {e}")
            # on tente de préserver ce qui a déjà été produit avant l'erreur
            try:
                if docx is not None: docx.save(docx_out)
            except Exception: pass
            try:
                if final_pdf is not None: final_pdf.save(pdf_out)
            except Exception: pass
            self._q_put("error", ("Erreur OCR", str(e)))
        finally:
            if txt_file is not None:
                try: txt_file.close()
                except Exception: pass


def main():
    app = OCRApp(); app.mainloop()


if __name__ == "__main__":
    main()
